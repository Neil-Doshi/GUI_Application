'''
The Purpose of this Code is to extract the content from a pdf file which consists of table of content for a project
Since the extracted text is not proper format i.e. same as table of content we use the code below to reformat it
For now this file only reads table of cable and generated word file for each entry in table of content.
'''


import re
import sys
import fitz
from docx import Document
from docx2pdf import convert
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING


class Submittal():
    def __init__(self, file, path, path_doc):
        self.document = fitz.open(file)  # table of content pdf file
        print(self.document)
        self.save_path = path
        self.path_doc = path_doc
        self.extracted_text_from_PDF = []  # variable holds the values read from pdf
        self.edge_case = []  # handles page end
        self.words = []
        self.indexs = []  # will hold all the numbers for each entry
        # Example : "4.9 Machine learning" is the entry in Table of content the list will hold 4.9
        self.relation = {}
        print("Code Initialized")
        self.pdf_text_extractor()

    def pdf_text_extractor(self):
        for page in self.document:
            words = page.get_text("words", sort=True) # iterates through each page and extracts every word in a list
            self.words += [word for word in words]
            self.extracted_text_from_PDF += [word[4] for word in words]  # saving the 4th value as that is the actual word we need
        self.edge_case += [i for i in range(len(self.words) - 1) if self.words[i][-3:] > self.words[i + 1][-3:]]
        if len(self.document) > len(self.edge_case):
            self.edge_case.append(len(self.words) - 1)
        print("Extracted text")
        self.index_extractor()

    def index_extractor(self):
        self.indexs = [i for i in range(len(self.extracted_text_from_PDF)) if re.search("^[1-9][.].*?[0-9]$", self.extracted_text_from_PDF[i])]  # iterating through all words and looking for number that resembles format 6.1..
        self.indexs.append(len(self.extracted_text_from_PDF))  # saving the value to handle last index of each page
        print("Extracted index")
        self.create_relation()

    def create_relation(self):
        helper_var1, helper_var2 = 0, 1  # helps with iteration of indexs
        while len(self.indexs) > 1:
            helper_string = ""  # helps hold value while combining the words
            for i in range(self.indexs[helper_var1] + 1, self.indexs[helper_var2]):  # we iterate through indexs as the values between 2 indexs is the entire entry from table of content
                if i in self.edge_case:  # handles edge case for last index of each page
                    helper_string += self.extracted_text_from_PDF[i]  # we end up adding 2 extra elements so we remove it
                    break
                helper_string += self.extracted_text_from_PDF[i] + " "  # adding all values to get the real string
            self.relation[self.extracted_text_from_PDF[self.indexs[helper_var1]]] = helper_string  # adding relation
            self.indexs = self.indexs[1:]  # removing first element for next iteration
        print("Generated Relations")
        self.create_doc()

    def create_doc(self):
        dummy = "0"
        for j in self.relation.keys():  #iterate through keys to combine and write to word file
            document = Document("C:\\Users\\Neil\\Downloads\\GUI_Application\\1.0.docx")  # initializes a blank document in memory
            if document.paragraphs[0]._element != None:  #lines 67 to 70 remove first paragraph
                p = document.paragraphs[0]._element
                p.getparent().remove(p)
                p._p = p._element = None
            if int(j[0]) > int(dummy):  #handles [1-9].0 case and makes a small table of content
                dummy = j[0]
                helper = self.case_handler(j[0])
                count = 0
                for k in helper.keys():
                    if count == 0:
                        para = document.add_paragraph(k + "    " + self.relation[k])
                        count += 1
                    else:
                        para = document.add_paragraph("\t" + k + "    " + self.relation[k])
            else:
                para = document.add_paragraph(j + "    " + self.relation[j])  # writes the key and value pair
            para.style.font.name = "Arial"
            para.style.font.size = Pt(20)
            para.style.paragraph_format.space_before = Pt(0.25)
            para.style.paragraph_format.space_after = Pt(0.25)
            n = str(j) + ".docx"  # Makes the files based on key
            document.save(self.path_doc + n)  # actually generates and save the files
        print("Generated word files")
        # self.create_pdf()
        
    def case_handler(self,pointer):
        res = {key: val for key, val in self.relation.items() if key.startswith(pointer)}
        return res
        

    def create_pdf(self):
        sys.stderr = open("consoleoutput.log", "w")
        convert(self.path_doc, self.save_path)
        return


if __name__ == '__main__':
    path = "C://Users//Neil//Downloads//GUI_Application//Result//"
    path_doc = "C://Users//Neil//Downloads//GUI_Application//Result//Doc//"
    file = "C://Users//Neil//Downloads//GUI_Application/Test_case//0.0.0.pdf"
    obj = Submittal(file, path, path_doc)
